import tempfile
from datetime import datetime
from typing import List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from app.schemas.meeting import ExportDocxRequest

SECTION_LABELS = {
    "en": {
        "title": "Meeting Summary",
        "language_label": "Language",
        "generated_label": "Generated",
        "summary": "Summary",
        "participants": "Participants",
        "decisions": "Decisions",
        "actionItems": "Action Items",
        "transcript": "Transcript",
        "empty": "N/A",
        "language_names": {"en": "English", "he": "Hebrew"},
    },
    "he": {
        "title": "סיכום פגישה",
        "language_label": "שפה",
        "generated_label": "נוצר בתאריך",
        "summary": "תקציר",
        "participants": "משתתפים",
        "decisions": "החלטות",
        "actionItems": "משימות לביצוע",
        "transcript": "תמלול מלא",
        "empty": "אין מידע",
        "language_names": {"en": "אנגלית", "he": "עברית"},
    },
}


def _is_rtl(language: str) -> bool:
    return language == "he"


def _set_font(style, font_name: str) -> None:
    style.font.name = font_name
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), font_name)


def _configure_document(doc: Document, rtl: bool) -> None:
    base_font = "Arial" if rtl else "Calibri"
    for style_name in ("Normal", "Heading 1", "Heading 2"):
        if style_name in doc.styles:
            _set_font(doc.styles[style_name], base_font)
    doc.styles["Normal"].font.size = Pt(11)


def _align_paragraph(paragraph, rtl: bool) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT


def _add_heading(doc: Document, text: str, level: int, rtl: bool):
    paragraph = doc.add_heading(text, level=level)
    _align_paragraph(paragraph, rtl)
    return paragraph


def _add_paragraph(doc: Document, text: str, rtl: bool, style: str | None = None):
    paragraph = doc.add_paragraph(text, style=style) if style else doc.add_paragraph(text)
    _align_paragraph(paragraph, rtl)
    return paragraph


def _add_section(
    doc: Document, title: str, items: List[str], rtl: bool, empty_label: str
) -> None:
    _add_heading(doc, title, level=2, rtl=rtl)
    if not items:
        _add_paragraph(doc, empty_label, rtl=rtl)
        return
    for item in items:
        _add_paragraph(doc, item, rtl=rtl, style="List Bullet")


def build_docx(payload: ExportDocxRequest) -> str:
    labels = SECTION_LABELS.get(payload.language, SECTION_LABELS["en"])
    rtl = _is_rtl(payload.language)
    doc = Document()
    _configure_document(doc, rtl)

    _add_heading(doc, labels["title"], level=1, rtl=rtl)
    language_name = labels["language_names"].get(payload.language, payload.language)
    _add_paragraph(doc, f'{labels["language_label"]}: {language_name}', rtl=rtl)
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
    _add_paragraph(doc, f'{labels["generated_label"]}: {timestamp}', rtl=rtl)

    _add_heading(doc, labels["summary"], level=2, rtl=rtl)
    _add_paragraph(doc, payload.summary or labels["empty"], rtl=rtl)

    _add_section(doc, labels["participants"], payload.participants, rtl, labels["empty"])
    _add_section(doc, labels["decisions"], payload.decisions, rtl, labels["empty"])
    _add_section(doc, labels["actionItems"], payload.actionItems, rtl, labels["empty"])

    _add_heading(doc, labels["transcript"], level=2, rtl=rtl)
    _add_paragraph(doc, payload.transcript or labels["empty"], rtl=rtl)

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    temp_file.close()
    doc.save(temp_file.name)
    return temp_file.name
